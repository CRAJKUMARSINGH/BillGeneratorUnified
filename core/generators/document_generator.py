"""
Document Generator - Generate billing documents from processed Excel data
"""
import pandas as pd
from typing import Dict, Any
import io
import asyncio
from functools import lru_cache
from jinja2 import Environment, FileSystemLoader
import os
import tempfile
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import re

class DocumentGenerator:
    """Generates various billing documents from processed Excel data using Jinja2 templates"""
    
    def __init__(self, data: Dict[str, Any]):
        self.data = data
        self.title_data = data.get('title_data', {})
        self.work_order_data = data.get('work_order_data', pd.DataFrame())
        self.bill_quantity_data = data.get('bill_quantity_data', pd.DataFrame())
        self.extra_items_data = data.get('extra_items_data', pd.DataFrame())
        
        # Set up Jinja2 environment for templates
        # template_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates')
        template_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'templates')
        self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
        
        # Prepare data for templates
        self.template_data = self._prepare_template_data()
    
    def _safe_float(self, value) -> float:
        """Safely convert value to float"""
        if pd.isna(value) or value is None or value == '':
            return 0.0
        try:
            return float(value)
        except (ValueError, TypeError):
            return 0.0
    
    def _safe_serial_no(self, value) -> str:
        """Safely convert serial number to string"""
        if pd.isna(value) or value is None:
            return ''
        return str(value)
    
    def _format_unit_or_text(self, value) -> str:
        """Format unit or text value"""
        if pd.isna(value) or value is None:
            return ''
        return str(value)
    
    def _format_number(self, value) -> str:
        """Format number for display"""
        if value == 0:
            return ''
        return f"{value:.2f}"
    
    def _number_to_words(self, num: int) -> str:
        """Convert number to words (simplified version)"""
        ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
        teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
        
        if num == 0:
            return 'Zero'
        
        if num < 10:
            return ones[num]
        elif num < 20:
            return teens[num - 10]
        elif num < 100:
            return tens[num // 10] + ('' if num % 10 == 0 else ' ' + ones[num % 10])
        elif num < 1000:
            return ones[num // 100] + ' Hundred' + ('' if num % 100 == 0 else ' ' + self._number_to_words(num % 100))
        elif num < 100000:
            return self._number_to_words(num // 1000) + ' Thousand' + ('' if num % 1000 == 0 else ' ' + self._number_to_words(num % 1000))
        else:
            return str(num)  # For very large numbers, just return as string
    
    def _has_extra_items(self) -> bool:
        """Check if there are extra items to include"""
        if isinstance(self.extra_items_data, pd.DataFrame):
            return not self.extra_items_data.empty
        return False
    
    def _prepare_template_data(self) -> Dict[str, Any]:
        """Prepare data structure for Jinja2 templates with enhanced first 20 rows handling"""
        # Calculate totals and prepare structured data
        work_items = []
        total_amount = 0
        
        # Process work order data
        for index, row in self.work_order_data.iterrows():
            quantity_since = self._safe_float(row.get('Quantity Since', row.get('Quantity', 0)))
            rate = self._safe_float(row.get('Rate', 0))
            amount = quantity_since * rate
            total_amount += amount
            
            work_items.append({
                'unit': row.get('Unit', ''),
                'quantity_since': quantity_since,
                'quantity_upto': self._safe_float(row.get('Quantity Upto', quantity_since)),
                'item_no': self._safe_serial_no(row.get('Item No.', row.get('Item', ''))),
                'description': row.get('Description', ''),
                'rate': rate,
                'amount_upto': amount,
                'amount_since': amount,
                'remark': row.get('Remark', '')
            })
        
        # Process extra items
        extra_items = []
        extra_total = 0
        
        if isinstance(self.extra_items_data, pd.DataFrame) and not self.extra_items_data.empty:
            for index, row in self.extra_items_data.iterrows():
                quantity = self._safe_float(row.get('Quantity', 0))
                rate = self._safe_float(row.get('Rate', 0))
                amount = quantity * rate
                extra_total += amount
                
                extra_items.append({
                    'unit': row.get('Unit', ''),
                    'quantity': quantity,
                    'item_no': self._safe_serial_no(row.get('Item No.', row.get('Item', ''))),
                    'description': row.get('Description', ''),
                    'rate': rate,
                    'amount': amount,
                    'remark': row.get('Remark', '')
                })
        
        # Calculate premiums
        tender_premium_percent = self._safe_float(self.title_data.get('TENDER PREMIUM %', 0))
        premium_amount = total_amount * (tender_premium_percent / 100)
        grand_total = total_amount + premium_amount
        
        extra_premium = extra_total * (tender_premium_percent / 100)
        extra_grand_total = extra_total + extra_premium
        
        # Calculate deductions and final amounts
        liquidated_damages = self._safe_float(self.title_data.get('Liquidated Damages', 0))
        sd_amount = grand_total * 0.10  # Security Deposit 10%
        it_amount = grand_total * 0.02  # Income Tax 2%
        gst_amount = grand_total * 0.02  # GST 2%
        lc_amount = grand_total * 0.01  # Labour Cess 1%
        total_deductions = sd_amount + it_amount + gst_amount + lc_amount + liquidated_damages
        
        # Get last bill amount
        last_bill_amount = self._safe_float(self.title_data.get('Amount Paid Vide Last Bill', 
                                                                 self.title_data.get('amount_paid_last_bill', 0)))
        net_payable = grand_total - last_bill_amount
        
        # Calculate totals data structure
        totals = {
            'grand_total': grand_total,
            'work_order_amount': total_amount,
            'tender_premium_percent': tender_premium_percent / 100,
            'tender_premium_amount': premium_amount,
            'final_total': grand_total,
            'extra_items_sum': extra_grand_total,
            'sd_amount': sd_amount,
            'it_amount': it_amount,
            'gst_amount': gst_amount,
            'lc_amount': lc_amount,
            'liquidated_damages': liquidated_damages,
            'total_deductions': total_deductions,
            'last_bill_amount': last_bill_amount,
            'net_payable': net_payable,
            'payable': grand_total,  # For compatibility
            'excess_amount': 0,  # Will be calculated from deviation data
            'excess_premium': 0,
            'excess_total': 0,
            'saving_amount': 0,
            'saving_premium': 0,
            'saving_total': 0,
            'net_difference': 0,
            'premium': {
                'percent': tender_premium_percent / 100 if tender_premium_percent > 0 else 0,
                'amount': premium_amount
            },
            'payable': grand_total,
            'liquidated_damages': self._safe_float(self.title_data.get('Liquidated Damages', 0))
        }
        
        # Prepare items data structure for templates
        items = []
        for item in work_items:
            # Skip rows with all zeros (empty rows)
            qty_since = item.get('quantity_since', 0) or 0
            qty_upto = item.get('quantity_upto', 0) or 0
            rate = item.get('rate', 0) or 0
            amount_upto = item.get('amount_upto', 0) or 0
            amount_since = item.get('amount_since', 0) or 0
            description = str(item.get('description', '')).strip()
            
            # Only add if there's actual data (non-zero values or description)
            if qty_since != 0 or qty_upto != 0 or rate != 0 or amount_upto != 0 or amount_since != 0 or description:
                items.append({
                    'unit': item['unit'],
                    'quantity_since_last': qty_since,
                    'quantity_upto_date': qty_upto,
                    'serial_no': item['item_no'],
                    'description': description,
                    'rate': rate,
                    'amount': amount_upto,
                    'amount_previous': amount_since,
                    'remark': item['remark']
                })
        
        # Add extra items to the items list
        for item in extra_items:
            qty = item.get('quantity', 0) or 0
            rate = item.get('rate', 0) or 0
            amount = item.get('amount', 0) or 0
            description = str(item.get('description', '')).strip()
            
            # Only add if there's actual data
            if qty != 0 or rate != 0 or amount != 0 or description:
                items.append({
                    'unit': item['unit'],
                    'quantity_since_last': qty,
                    'quantity_upto_date': qty,
                    'serial_no': item['item_no'],
                    'description': description,
                    'rate': rate,
                    'amount': amount,
                    'amount_previous': amount,
                    'remark': item['remark']
                })
        
        # Prepare deviation data
        deviation_items = []
        for item in work_items:
            deviation_items.append({
                'serial_no': item['item_no'],
                'description': item['description'],
                'unit': item['unit'],
                'qty_wo': item['quantity_since'],
                'rate': item['rate'],
                'amt_wo': item['amount_since'],
                'qty_bill': item['quantity_upto'],
                'amt_bill': item['amount_upto'],
                'excess_qty': max(0, item['quantity_upto'] - item['quantity_since']),
                'excess_amt': max(0, item['amount_upto'] - item['amount_since']),
                'saving_qty': max(0, item['quantity_since'] - item['quantity_upto']),
                'saving_amt': max(0, item['amount_since'] - item['amount_upto']),
                'remark': item['remark']
            })
        
        # Prepare summary data
        summary = {
            'work_order_total': total_amount,
            'executed_total': total_amount,
            'overall_excess': 0,
            'overall_saving': 0,
            'tender_premium_f': premium_amount,
            'tender_premium_h': premium_amount,
            'tender_premium_j': 0,
            'tender_premium_l': 0,
            'grand_total_f': grand_total,
            'grand_total_h': grand_total,
            'grand_total_j': 0,
            'grand_total_l': 0,
            'net_difference': 0,
            'percentage_deviation': 0,
            'is_saving': False,
            'premium': {
                'percent': tender_premium_percent / 100 if tender_premium_percent > 0 else 0,
                'amount': premium_amount
            }
        }
        
        # Ensure all data is available for templates with enhanced first 20 rows handling
        template_data = {
            'title_data': self.title_data,
            'work_items': work_items,
            'extra_items': extra_items,
            'totals': totals,
            'current_date': datetime.now().strftime('%d/%m/%Y'),
            'total_amount': total_amount,
            'tender_premium_percent': tender_premium_percent,
            'premium_amount': premium_amount,
            'grand_total': grand_total,
            'extra_total': extra_total,
            'extra_premium': extra_premium,
            'extra_grand_total': extra_grand_total,
            'final_total': grand_total + extra_grand_total,
            'payable_words': self._number_to_words(int(net_payable)),
            'notes': ['Work completed as per schedule', 'All measurements verified', 'Quality as per specifications'],
            'items': items,
            'deviation_items': deviation_items,
            'summary': summary,
            'agreement_no': self.title_data.get('Work Order No', ''),
            'name_of_work': self.title_data.get('Project Name', ''),
            'name_of_firm': self.title_data.get('Contractor Name', ''),
            'date_commencement': self.title_data.get('Date of Commencement', ''),
            'date_completion': self.title_data.get('Date of Completion', ''),
            'actual_completion': self.title_data.get('Actual Date of Completion', ''),
            'work_order_amount': total_amount,
            'bill_grand_total': grand_total,
            'extra_items_sum': extra_grand_total,
            'delay_days': 0,
            'header': [],
            'measurement_officer': self.title_data.get('Measurement Officer', 'Junior Engineer'),
            'measurement_date': self.title_data.get('Measurement Date', datetime.now().strftime('%d/%m/%Y')),
            'measurement_book_page': self.title_data.get('Measurement Book Page', '04-20'),
            'measurement_book_no': self.title_data.get('Measurement Book No', '887'),
            'officer_name': self.title_data.get('Officer Name', 'Name of Officer'),
            'officer_designation': self.title_data.get('Officer Designation', 'Assistant Engineer'),
            'bill_date': self.title_data.get('Bill Date', '__/__/____'),
            'authorising_officer_name': self.title_data.get('Authorising Officer Name', 'Name of Authorising Officer'),
            'authorising_officer_designation': self.title_data.get('Authorising Officer Designation', 'Executive Engineer'),
            'authorisation_date': self.title_data.get('Authorisation Date', '__/__/____'),
            'payable_amount': net_payable,
            'amount_words': self._number_to_words(int(net_payable)),
            # Enhanced first 20 rows data for validation
            'first_20_rows_processed': self.title_data.get('_first_20_rows_processed', False),
            'first_20_rows_count': self.title_data.get('_first_20_rows_count', 0)
        }
        
        # No title image data URI needed - using direct HTML title instead
        
        # Add header data if available
        if 'header' in self.title_data:
            template_data['header'] = self.title_data['header']
        
        return template_data
    
    def generate_all_documents(self) -> Dict[str, str]:
        """
        Generate all required documents using Jinja2 templates
        
        Returns:
            Dictionary containing all generated documents in HTML format
        """
        documents = {}
        
        try:
            # Generate individual documents using templates
            documents['First Page Summary'] = self._render_template('first_page.html')
            documents['Deviation Statement'] = self._render_template('deviation_statement.html') 
            documents['BILL SCRUTINY SHEET'] = self._render_template('note_sheet.html')
            
            # Only generate Extra Items document if there are extra items
            if self._has_extra_items():
                documents['Extra Items Statement'] = self._render_template('extra_items.html')
            
            documents['Certificate II'] = self._render_template('certificate_ii.html')
            documents['Certificate III'] = self._render_template('certificate_iii.html')
        except Exception as e:
            print(f"Template rendering failed, falling back to programmatic generation: {e}")
            # Fallback to programmatic generation if templates fail
            documents['First Page Summary'] = self._generate_first_page()
            documents['Deviation Statement'] = self._generate_deviation_statement()
            documents['BILL SCRUTINY SHEET'] = self._generate_final_bill_scrutiny()
            
            # Only generate Extra Items document if there are extra items
            if self._has_extra_items():
                documents['Extra Items Statement'] = self._generate_extra_items_statement()
            
            documents['Certificate II'] = self._generate_certificate_ii()
            documents['Certificate III'] = self._generate_certificate_iii()
        
        return documents
    
    def generate_doc_documents(self) -> Dict[str, bytes]:
        """
        Generate all required documents in DOC format
        
        Returns:
            Dictionary containing all generated documents in DOC format (bytes)
        """
        doc_documents = {}
        
        # Generate DOC versions of all documents
        doc_documents['First Page Summary.docx'] = self._generate_doc_first_page()
        doc_documents['Deviation Statement.docx'] = self._generate_doc_deviation_statement()
        doc_documents['BILL SCRUTINY SHEET.docx'] = self._generate_doc_note_sheet()
        
        # Only generate Extra Items document if there are extra items
        if self._has_extra_items():
            doc_documents['Extra Items Statement.docx'] = self._generate_doc_extra_items()
        
        doc_documents['Certificate II.docx'] = self._generate_doc_certificate_ii()
        doc_documents['Certificate III.docx'] = self._generate_doc_certificate_iii()
        
        return doc_documents
    
    def _render_template(self, template_name: str) -> str:
        """Render a Jinja2 template with the prepared data"""
        try:
            template = self.jinja_env.get_template(template_name)
            # Pass both the template data and the original data to the template
            render_data = {'data': self.template_data}
            render_data.update(self.template_data)
            return template.render(**render_data)
        except Exception as e:
            print(f"Failed to render template {template_name}: {e}")
            raise
    
    async def _convert_html_to_pdf_async(self, html_content: str, doc_name: str) -> bytes:
        """
        Convert HTML to PDF using Playwright for pixel-perfect rendering
        PERMANENT FIX: Tables will NOT shrink
        Deviation Statement uses LANDSCAPE orientation
        """
        from playwright.async_api import async_playwright
        
        # Determine page orientation based on document name
        is_landscape = 'Deviation' in doc_name
        page_size = 'A4 landscape' if is_landscape else 'A4'
        
        # Add CSS to prevent table shrinking + CSS zoom for pixel-perfect rendering
        no_shrink_css = f"""
        <style>
            /* CRITICAL: CSS Zoom for pixel-perfect scaling */
            body {{
                zoom: 1.0;
                -moz-transform: scale(1.0);
                -moz-transform-origin: 0 0;
            }}
            
            /* CRITICAL: Prevent table shrinking */
            table {{
                table-layout: fixed !important;
                width: 100% !important;
                min-width: 100% !important;
                max-width: 100% !important;
                border-collapse: collapse !important;
            }}
            th, td {{
                white-space: normal !important;
                word-wrap: break-word !important;
                overflow-wrap: break-word !important;
                box-sizing: border-box !important;
            }}
            
            /* Disable text rendering optimizations */
            * {{
                -webkit-font-smoothing: antialiased !important;
                -moz-osx-font-smoothing: grayscale !important;
                text-rendering: optimizeLegibility !important;
            }}
        </style>
        """
        
        # Inject CSS into HTML
        if '<head>' in html_content:
            html_content = html_content.replace('<head>', f'<head>\n{no_shrink_css}', 1)
        else:
            html_content = html_content.replace('<html>', f'<html>\n<head>\n{no_shrink_css}\n</head>', 1)
        
        # Launch Playwright and convert to PDF
        async with async_playwright() as p:
            # Use Chromium browser
            browser = await p.chromium.launch(
                args=[
                    '--no-sandbox',
                    '--disable-setuid-sandbox',
                    '--disable-dev-shm-usage',
                    '--disable-accelerated-2d-canvas',
                    '--no-first-run',
                    '--no-zygote',
                    '--single-process'  # Critical for Streamlit Cloud compatibility
                ]
            )
            page = await browser.new_page()
            
            # Set viewport to A4 size at 96 DPI
            await page.set_viewport_size({'width': 794, 'height': 1123})
            
            # Set the HTML content
            await page.set_content(html_content, wait_until='networkidle')
            
            # Wait for rendering
            await page.wait_for_timeout(500)
            
            # Generate PDF with NO SHRINKING settings
            # Use landscape for Deviation Statement
            pdf_bytes = await page.pdf(
                format='A4',
                landscape=is_landscape,  # Landscape for Deviation Statement
                print_background=True,
                margin={'top': '0mm', 'right': '0mm', 'bottom': '0mm', 'left': '0mm'},
                prefer_css_page_size=True,
                display_header_footer=False,
                scale=1.0  # CRITICAL: No scaling = no shrinking
            )
            
            await browser.close()
            return pdf_bytes
    
    def create_pdf_documents(self, documents: Dict[str, str]) -> Dict[str, bytes]:
        """
        Convert HTML documents to PDF format using Playwright for exact HTML rendering
        
        Args:
            documents: Dictionary of HTML documents
            
        Returns:
            Dictionary of PDF documents as bytes
        """
        pdf_documents = {}
        
        # Convert each document to PDF asynchronously
        async def convert_all():
            tasks = []
            for doc_name, html_content in documents.items():
                task = self._convert_html_to_pdf_async(html_content, doc_name)
                tasks.append(task)
            
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            for i, (doc_name, result) in enumerate(zip(documents.keys(), results)):
                if isinstance(result, Exception):
                    print(f"Failed to convert {doc_name} to PDF: {result}")
                    # Fallback to xhtml2pdf if Playwright fails
                    try:
                        pdf_documents[doc_name] = self._convert_html_to_pdf_fallback(documents[doc_name], doc_name)
                    except Exception as fallback_error:
                        print(f"Fallback PDF conversion also failed for {doc_name}: {fallback_error}")
                else:
                    pdf_documents[doc_name] = result
        
        # Run the async conversion
        asyncio.run(convert_all())
        
        return pdf_documents
    
    def _convert_html_to_pdf_fallback(self, html_content: str, doc_name: str) -> bytes:
        """
        Fallback PDF conversion using xhtml2pdf (works on Streamlit Cloud)
        """
        try:
            print("[INFO] Using xhtml2pdf (Streamlit Cloud compatible)...")
            from xhtml2pdf import pisa
            import io
            
            # Add CSS zoom to HTML for better rendering
            zoom = 1.0
            html_with_zoom = self.add_css_zoom_to_html(html_content, zoom)
            
            output = io.BytesIO()
            pisa.CreatePDF(html_with_zoom, dest=output, encoding="utf-8")
            return output.getvalue()
        except Exception as e:
            print(f"[WARNING] xhtml2pdf failed: {e}")
            raise Exception("No PDF engine available. Please check requirements installation.")
    
    def add_css_zoom_to_html(self, html_content: str, zoom: float = 1.0) -> str:
        """
        Add CSS zoom to HTML content for better rendering
        """
        zoom_style = f"""
        <style>
            body {{
                zoom: {zoom};
                -moz-transform: scale({zoom});
                -moz-transform-origin: 0 0;
            }}
        </style>
        """
        
        if '<head>' in html_content:
            return html_content.replace('<head>', f'<head>\n{zoom_style}', 1)
        else:
            return html_content.replace('<html>', f'<html>\n<head>\n{zoom_style}\n</head>', 1)
    
    def calculate_optimal_zoom(self, content_width_px: int, page_width_mm: float = 210) -> float:
        """
        Calculate optimal zoom level for pixel-perfect fit
        
        Args:
            content_width_px: Content width in pixels
            page_width_mm: Page width in millimeters (default A4 = 210mm)
            
        Returns:
            Optimal zoom level
        """
        # Convert mm to pixels at 96 DPI
        page_width_px = (page_width_mm / 25.4) * 96
        
        # Calculate zoom to fit content
        zoom = page_width_px / content_width_px
        
        # Round to 2 decimal places
        return round(zoom, 2)
    
    def batch_convert(self, html_documents: Dict[str, str], 
                     output_dir: str = "output_pdfs",
                     enable_fallback: bool = True) -> Dict[str, str]:
        """
        Batch convert HTML documents to PDFs
        
        Args:
            html_documents: Dictionary mapping document names to HTML content
            output_dir: Directory to save PDFs
            enable_fallback: Enable fallback to xhtml2pdf if Playwright fails
            
        Returns:
            Dictionary mapping document names to PDF file paths
        """
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Convert documents to PDFs
        pdf_paths = {}
        pdf_documents = self.create_pdf_documents(html_documents)
        
        # Save PDFs to files
        for doc_name, pdf_content in pdf_documents.items():
            # Generate safe filename
            safe_name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', doc_name)
            if not safe_name.endswith('.pdf'):
                safe_name += '.pdf'
            
            pdf_path = os.path.join(output_dir, safe_name)
            
            try:
                with open(pdf_path, 'wb') as f:
                    f.write(pdf_content)
                pdf_paths[doc_name] = pdf_path
                print(f"Saved {doc_name} to {pdf_path}")
            except Exception as e:
                print(f"Failed to save {doc_name}: {e}")
        
        return pdf_paths
    
    def _generate_first_page(self) -> str:
        """Generate First Page Summary document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>First Page Summary</title>
            <style>
                @page {{ 
                    size: A4; 
                    margin: 10mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 10mm;
                    font-size: 10pt; 
                }}
                .header {{ text-align: center; margin-bottom: 8px; }}
                .subtitle {{ font-size: 11pt; margin: 3px 0; }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 6px 0; 
                    table-layout: fixed;
                }}
                thead {{ display: table-header-group; }}
                tr, img {{ break-inside: avoid; }}
                th, td {{ 
                    border: 1px solid #000; 
                    padding: 4px; 
                    text-align: left; 
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    vertical-align: top;
                }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
                .total-row {{ font-weight: bold; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="subtitle">First Page Summary</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <h3>Project Information</h3>
            <table>
                <tr><td><strong>Project Name:</strong></td><td>{self.title_data.get('Project Name', 'N/A')}</td></tr>
                <tr><td><strong>Contract No:</strong></td><td>{self.title_data.get('Contract No', 'N/A')}</td></tr>
                <tr><td><strong>Work Order No:</strong></td><td>{self.title_data.get('Work Order No', 'N/A')}</td></tr>
            </table>
            
            <h3>Work Items Summary</h3>
            <table>
                <thead>
                    <tr>
                        <th style="width: 10.06mm;">Unit</th>
                        <th style="width: 13.76mm;">Quantity executed (or supplied) since last certificate</th>
                        <th style="width: 13.76mm;">Quantity executed (or supplied) upto date as per MB</th>
                        <th style="width: 9.55mm;">S. No.</th>
                        <th style="width: 63.83mm;">Item of Work supplies (Grouped under "sub-head" and "sub work" of estimate)</th>
                        <th style="width: 13.16mm;">Rate</th>
                        <th style="width: 19.53mm;">Upto date Amount</th>
                        <th style="width: 15.15mm;">Amount Since previous bill (Total for each sub-head)</th>
                        <th style="width: 11.96mm;">Remarks</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        # Get tender premium percentage from title data first (needed for calculations)
        tender_premium_percent = self._safe_float(self.title_data.get('TENDER PREMIUM %', 
                                                                       self.title_data.get('Tender Premium %', 0)))
        
        # Add work order items
        total_amount = 0
        running_total = 0
        
        for index, row in self.work_order_data.iterrows():
            description = str(row.get('Description', '')).strip()
            
            # Check if this is a summary row (Total, Add Tender Premium, Grand Total)
            is_total_row = description.lower() == 'total'
            is_premium_row = 'premium' in description.lower() and 'add' in description.lower()
            is_grand_total_row = 'grand' in description.lower() and 'total' in description.lower()
            
            if is_total_row:
                # This is the TOTAL row - show running total
                qty_since_display = ""
                qty_upto_display = ""
                rate_display = ""
                amt_upto_display = f"{running_total:.2f}"
                amt_since_display = f"{running_total:.2f}"
                total_amount = running_total
            elif is_premium_row:
                # This is the TENDER PREMIUM row
                premium_amount = running_total * (tender_premium_percent / 100)
                qty_since_display = ""
                qty_upto_display = ""
                rate_display = f"{tender_premium_percent:.2f}%"
                amt_upto_display = f"{premium_amount:.2f}"
                amt_since_display = f"{premium_amount:.2f}"
                total_amount += premium_amount
                running_total += premium_amount
            elif is_grand_total_row:
                # This is the GRAND TOTAL row
                qty_since_display = ""
                qty_upto_display = ""
                rate_display = ""
                amt_upto_display = f"{running_total:.2f}"
                amt_since_display = f"{running_total:.2f}"
            else:
                # Regular work item
                qty_since = self._safe_float(row.get('Quantity Since', row.get('Quantity', 0)))
                qty_upto = self._safe_float(row.get('Quantity Upto', qty_since))
                rate = self._safe_float(row.get('Rate', 0))
                amt_upto = qty_upto * rate
                amt_since = qty_since * rate
                
                # Update running totals
                running_total += amt_upto
                total_amount += amt_upto
                
                # Format for display (empty string if zero)
                qty_since_display = f"{qty_since:.2f}" if qty_since != 0 else ""
                qty_upto_display = f"{qty_upto:.2f}" if qty_upto != 0 else ""
                rate_display = f"{rate:.2f}" if rate != 0 else ""
                amt_upto_display = f"{amt_upto:.2f}" if amt_upto != 0 else ""
                amt_since_display = f"{amt_since:.2f}" if amt_since != 0 else ""
            
            html_content += f"""
                    <tr>
                        <td>{row.get('Unit', '')}</td>
                        <td class="amount">{qty_since_display}</td>
                        <td class="amount">{qty_upto_display}</td>
                        <td>{self._safe_serial_no(row.get('Item No.', row.get('Item', '')))}</td>
                        <td>{description}</td>
                        <td class="amount">{rate_display}</td>
                        <td class="amount">{amt_upto_display}</td>
                        <td class="amount">{amt_since_display}</td>
                        <td>{row.get('Remark', '')}</td>
                    </tr>
            """
        
        html_content += f"""
                </tbody>
            </table>
            
            <!-- Closure Lines -->
            <div style="margin-top: 40px; width: 100%; display: table;">
                <div style="display: table-row;">
                    <div style="display: table-cell; width: 50%; padding: 10px; vertical-align: top;">
                        <div style="border-top: 2px solid #000; padding-top: 5px; text-align: center;">
                            <strong>Prepared by</strong><br>
                            <span style="font-size: 8pt;">Assistant Engineer</span><br>
                            <span style="font-size: 8pt;">Date: {current_date}</span>
                        </div>
                    </div>
                    <div style="display: table-cell; width: 50%; padding: 10px; vertical-align: top;">
                        <div style="border-top: 2px solid #000; padding-top: 5px; text-align: center;">
                            <strong>Checked & Approved by</strong><br>
                            <span style="font-size: 8pt;">Executive Engineer</span><br>
                            <span style="font-size: 8pt;">Date: {current_date}</span>
                        </div>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_deviation_statement(self) -> str:
        """Generate Deviation Statement document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Deviation Statement</title>
            <style>
                @page {{ 
                    size: A4 landscape; 
                    margin: 10mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 10mm;
                    font-size: 9pt; 
                }}
                .header {{ text-align: center; margin-bottom: 8px; }}
                .subtitle {{ font-size: 10pt; margin: 3px 0; }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 4px 0; 
                    table-layout: fixed;
                }}
                thead {{ display: table-header-group; }}
                tr, img {{ break-inside: avoid; }}
                th, td {{ 
                    border: 1px solid #000; 
                    padding: 3px; 
                    text-align: left; 
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    font-size: 8.5pt;
                }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
                
                /* Column-specific widths for Deviation Statement (percentages) */
                th:nth-child(1), td:nth-child(1) {{ width: 2.20%; }}   /* Item No */
                th:nth-child(2), td:nth-child(2) {{ width: 43.22%; }}  /* Description */
                th:nth-child(3), td:nth-child(3) {{ width: 3.85%; }}   /* Unit */
                th:nth-child(4), td:nth-child(4) {{ width: 3.85%; }}   /* Qty WO */
                th:nth-child(5), td:nth-child(5) {{ width: 3.85%; }}   /* Rate */
                th:nth-child(6), td:nth-child(6) {{ width: 3.85%; }}   /* Amt WO */
                th:nth-child(7), td:nth-child(7) {{ width: 3.85%; }}   /* Qty Exec */
                th:nth-child(8), td:nth-child(8) {{ width: 3.85%; }}   /* Amt Exec */
                th:nth-child(9), td:nth-child(9) {{ width: 3.85%; }}   /* Excess Qty */
                th:nth-child(10), td:nth-child(10) {{ width: 3.85%; }} /* Excess Amt */
                th:nth-child(11), td:nth-child(11) {{ width: 3.85%; }} /* Saving Qty */
                th:nth-child(12), td:nth-child(12) {{ width: 3.85%; }} /* Saving Amt */
                th:nth-child(13), td:nth-child(13) {{ width: 17.58%; }} /* Remarks */
            </style>
        </head>
        <body>
            <div class="header">
                <div class="subtitle">Deviation Statement</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <table>
                <thead>
                    <tr>
                        <th style="width: 6mm;">ITEM No.</th>
                        <th style="width: 118mm;">Description</th>
                        <th style="width: 10.5mm;">Unit</th>
                        <th style="width: 10.5mm;">Qty as per Work Order</th>
                        <th style="width: 10.5mm;">Rate</th>
                        <th style="width: 10.5mm;">Amt as per Work Order Rs.</th>
                        <th style="width: 10.5mm;">Qty Executed</th>
                        <th style="width: 10.5mm;">Amt as per Executed Rs.</th>
                        <th style="width: 10.5mm;">Excess Qty</th>
                        <th style="width: 10.5mm;">Excess Amt Rs.</th>
                        <th style="width: 10.5mm;">Saving Qty</th>
                        <th style="width: 10.5mm;">Saving Amt Rs.</th>
                        <th style="width: 48mm;">REMARKS/ REASON.</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        # Compare work order with bill quantity data
        for index, wo_row in self.work_order_data.iterrows():
            bq_row = None
            if isinstance(self.bill_quantity_data, pd.DataFrame) and not self.bill_quantity_data.empty:
                wo_item = wo_row.get('Item No.', wo_row.get('Item', ''))
                bq_item_col = 'Item No.' if 'Item No.' in self.bill_quantity_data.columns else 'Item'
                matching_rows = self.bill_quantity_data[
                    self.bill_quantity_data[bq_item_col] == wo_item
                ]
                if isinstance(matching_rows, pd.DataFrame) and not matching_rows.empty:
                    bq_row = matching_rows.iloc[0]
            
            wo_qty = self._safe_float(wo_row.get('Quantity Since', wo_row.get('Quantity', 0)))
            wo_rate = self._safe_float(wo_row.get('Rate', 0))
            wo_amount = wo_qty * wo_rate
            
            bq_qty = self._safe_float(bq_row.get('Quantity', 0)) if bq_row is not None else 0
            bq_amount = bq_qty * wo_rate
            
            excess_qty = max(0, bq_qty - wo_qty)
            excess_amt = excess_qty * wo_rate
            saving_qty = max(0, wo_qty - bq_qty)
            saving_amt = saving_qty * wo_rate
            
            wo_qty_display = f"{wo_qty:.2f}" if wo_qty > 0 else ""
            wo_rate_display = f"{wo_rate:.2f}" if wo_rate > 0 else ""
            wo_amount_display = f"{wo_amount:.2f}" if wo_amount > 0 else ""
            bq_qty_display = f"{bq_qty:.2f}" if bq_qty > 0 else ""
            bq_amount_display = f"{bq_amount:.2f}" if bq_amount > 0 else ""
            excess_qty_display = f"{excess_qty:.2f}" if excess_qty > 0 else ""
            excess_amt_display = f"{excess_amt:.2f}" if excess_amt > 0 else ""
            saving_qty_display = f"{saving_qty:.2f}" if saving_qty > 0 else ""
            saving_amt_display = f"{saving_amt:.2f}" if saving_amt > 0 else ""
            
            html_content += f"""
                    <tr>
                        <td>{self._safe_serial_no(wo_row.get('Item No.', wo_row.get('Item', '')))}</td>
                        <td>{wo_row.get('Description', '')}</td>
                        <td>{wo_row.get('Unit', '')}</td>
                        <td class="amount">{wo_qty_display}</td>
                        <td class="amount">{wo_rate_display}</td>
                        <td class="amount">{wo_amount_display}</td>
                        <td class="amount">{bq_qty_display}</td>
                        <td class="amount">{bq_amount_display}</td>
                        <td class="amount">{excess_qty_display}</td>
                        <td class="amount">{excess_amt_display}</td>
                        <td class="amount">{saving_qty_display}</td>
                        <td class="amount">{saving_amt_display}</td>
                        <td>{wo_row.get('Remark', '')}</td>
                    </tr>
            """
        
        html_content += f"""
                </tbody>
            </table>
            
            <!-- Closure Lines -->
            <div style="margin-top: 40px; width: 100%; display: table;">
                <div style="display: table-row;">
                    <div style="display: table-cell; width: 50%; padding: 10px; vertical-align: top;">
                        <div style="border-top: 2px solid #000; padding-top: 5px; text-align: center;">
                            <strong>Prepared by</strong><br>
                            <span style="font-size: 8pt;">Assistant Engineer</span><br>
                            <span style="font-size: 8pt;">Date: {current_date}</span>
                        </div>
                    </div>
                    <div style="display: table-cell; width: 50%; padding: 10px; vertical-align: top;">
                        <div style="border-top: 2px solid #000; padding-top: 5px; text-align: center;">
                            <strong>Checked & Approved by</strong><br>
                            <span style="font-size: 8pt;">Executive Engineer</span><br>
                            <span style="font-size: 8pt;">Date: {current_date}</span>
                        </div>
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_final_bill_scrutiny(self) -> str:
        """Generate Bill Scrutiny Sheet"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        bill_number = self.title_data.get('Bill Number', 'Bill Number Not Available')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>BILL SCRUTINY SHEET</title>
            <style>
                @page {{ 
                    size: A4; 
                    margin: 10mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 10mm;
                    font-size: 10pt; 
                }}
                .header {{ text-align: center; margin-bottom: 8px; }}
                .subtitle {{ font-size: 11pt; margin: 3px 0; }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 6px 0; 
                    table-layout: fixed;
                }}
                thead {{ display: table-header-group; }}
                tr, img {{ break-inside: avoid; }}
                th, td {{ 
                    border: 1px solid #000; 
                    padding: 4px; 
                    text-align: left; 
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="subtitle">BILL SCRUTINY SHEET - {bill_number}</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <h3>Bill Summary</h3>
            <table>
                <thead>
                    <tr>
                        <th style="width: 10%;">Item No.</th>
                        <th style="width: 40%;">Description</th>
                        <th style="width: 10%;">Unit</th>
                        <th style="width: 13%;">Quantity</th>
                        <th style="width: 13%;">Rate</th>
                        <th style="width: 14%;">Amount</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        total_amount = 0
        if isinstance(self.bill_quantity_data, pd.DataFrame) and not self.bill_quantity_data.empty:
            for index, row in self.bill_quantity_data.iterrows():
                quantity = self._safe_float(row.get('Quantity', 0))
                rate = self._safe_float(row.get('Rate', 0))
                amount = quantity * rate
                total_amount += amount
                
                html_content += f"""
                        <tr>
                            <td>{self._safe_serial_no(row.get('Item No.', row.get('Item', '')))}</td>
                            <td>{row.get('Description', '')}</td>
                            <td>{self._format_unit_or_text(row.get('Unit', ''))}</td>
                            <td class="amount">{self._format_number(quantity)}</td>
                            <td class="amount">{self._format_number(rate)}</td>
                            <td class="amount">{self._format_number(amount)}</td>
                        </tr>
                """
        else:
            html_content += """
                    <tr>
                        <td colspan="6">No bill quantity data available</td>
                    </tr>
            """
        
        html_content += f"""
                    <tr style="font-weight: bold;">
                        <td colspan="5">TOTAL</td>
                        <td class="amount">{total_amount:.0f}</td>
                    </tr>
                </tbody>
            </table>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_extra_items_statement(self) -> str:
        """Generate Extra Items Statement"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Extra Items Statement</title>
            <style>
                @page {{ 
                    size: A4; 
                    margin: 10mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 10mm;
                    font-size: 10pt; 
                }}
                .header {{ text-align: center; margin-bottom: 8px; }}
                .subtitle {{ font-size: 11pt; margin: 3px 0; }}
                table {{ 
                    width: 100%; 
                    border-collapse: collapse; 
                    margin: 6px 0; 
                    table-layout: fixed;
                }}
                thead {{ display: table-header-group; }}
                tr, img {{ break-inside: avoid; }}
                th, td {{ 
                    border: 1px solid #000; 
                    padding: 4px; 
                    text-align: left; 
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .amount {{ text-align: right; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="subtitle">Extra Items Statement</div>
                <div class="subtitle">Date: {current_date}</div>
            </div>
            
            <h3>Extra Items</h3>
            <table>
                <thead>
                    <tr>
                        <th style="width: 10%;">Item No.</th>
                        <th style="width: 50%;">Description</th>
                        <th style="width: 10%;">Unit</th>
                        <th style="width: 10%;">Quantity</th>
                        <th style="width: 10%;">Rate</th>
                        <th style="width: 10%;">Amount</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        total_amount = 0
        if isinstance(self.extra_items_data, pd.DataFrame) and not self.extra_items_data.empty:
            for index, row in self.extra_items_data.iterrows():
                quantity = self._safe_float(row.get('Quantity', 0))
                rate = self._safe_float(row.get('Rate', 0))
                amount = quantity * rate
                total_amount += amount
                
                html_content += f"""
                        <tr>
                            <td>{self._safe_serial_no(row.get('Item No.', row.get('Item', '')))}</td>
                            <td>{row.get('Description', '')}</td>
                            <td>{self._format_unit_or_text(row.get('Unit', ''))}</td>
                            <td class="amount">{self._format_number(quantity)}</td>
                            <td class="amount">{self._format_number(rate)}</td>
                            <td class="amount">{self._format_number(amount)}</td>
                        </tr>
                """
        else:
            html_content += """
                    <tr>
                        <td colspan="6">No extra items data available</td>
                    </tr>
            """
        
        html_content += f"""
                    <tr style="font-weight: bold;">
                        <td colspan="5">TOTAL</td>
                        <td class="amount">{total_amount:.2f}</td>
                    </tr>
                </tbody>
            </table>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_certificate_ii(self) -> str:
        """Generate Certificate II document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Certificate II - Work Completion Certificate</title>
            <style>
                @page {{ 
                    size: A4; 
                    margin: 15mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 15mm;
                    font-size: 11pt; 
                    line-height: 1.6;
                }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 14pt; font-weight: bold; margin-bottom: 30px; }}
                .content {{ text-align: justify; margin-bottom: 30px; }}
                .signature-block {{ 
                    margin-top: 50px; 
                    display: flex; 
                    justify-content: space-between;
                }}
                .signature {{ 
                    width: 45%; 
                    text-align: center;
                }}
                .signature-line {{ 
                    margin-top: 60px; 
                    border-top: 1px solid #000; 
                    padding-top: 5px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">CERTIFICATE II</div>
                <div class="title">WORK COMPLETION CERTIFICATE</div>
            </div>
            
            <div class="content">
                <p>This is to certify that the work mentioned in the agreement/bill has been completed/executed as per the terms 
                and conditions of the agreement and measurement book. The quality of work is found satisfactory.</p>
                
                <p>The bills have been scrutinized and found correct and in order. The work has been executed as per specifications 
                and the measurements recorded are correct.</p>
                
                <p>All materials used are as per approved standards and specifications. The work has been carried out under my 
                personal supervision and knowledge.</p>
            </div>
            
            <div class="signature-block">
                <div class="signature">
                    <div>Prepared by</div>
                    <div class="signature-line">
                        Assistant Engineer<br>
                        Date: {current_date}
                    </div>
                </div>
                <div class="signature">
                    <div>Approved by</div>
                    <div class="signature-line">
                        Executive Engineer<br>
                        Date: {current_date}
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_certificate_iii(self) -> str:
        """Generate Certificate III document"""
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Certificate III - Payment Certification</title>
            <style>
                @page {{ 
                    size: A4; 
                    margin: 15mm;
                }}
                * {{ box-sizing: border-box; }}
                body {{ 
                    font-family: Arial, sans-serif; 
                    margin: 0; 
                    padding: 15mm;
                    font-size: 11pt; 
                    line-height: 1.6;
                }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .title {{ font-size: 14pt; font-weight: bold; margin-bottom: 30px; }}
                .content {{ text-align: justify; margin-bottom: 30px; }}
                .signature-block {{ 
                    margin-top: 50px; 
                    display: flex; 
                    justify-content: space-between;
                }}
                .signature {{ 
                    width: 45%; 
                    text-align: center;
                }}
                .signature-line {{ 
                    margin-top: 60px; 
                    border-top: 1px solid #000; 
                    padding-top: 5px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="title">CERTIFICATE III</div>
                <div class="title">PAYMENT CERTIFICATION</div>
            </div>
            
            <div class="content">
                <p>This is to certify that the payment claimed in the bill is correct and in order. The work has been executed 
                satisfactorily and measurements are correct as per records.</p>
                
                <p>The bills have been thoroughly examined and verified. All deductions as per rules have been made. The payment 
                is recommended subject to budget provision and availability of funds.</p>
                
                <p>The work has been completed/executed as per agreement terms and conditions. No adverse remarks are reported 
                against the contractor.</p>
            </div>
            
            <div class="signature-block">
                <div class="signature">
                    <div>Prepared by</div>
                    <div class="signature-line">
                        Assistant Engineer<br>
                        Date: {current_date}
                    </div>
                </div>
                <div class="signature">
                    <div>Approved by</div>
                    <div class="signature-line">
                        Executive Engineer<br>
                        Date: {current_date}
                    </div>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_doc_first_page(self) -> bytes:
        """Generate First Page Summary document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('First Page Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add project information
        doc.add_heading('Project Information', level=1)
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Populate project info
        table.cell(0, 0).text = 'Project Name:'
        table.cell(0, 1).text = str(self.title_data.get('Project Name', 'N/A'))
        table.cell(1, 0).text = 'Contract No:'
        table.cell(1, 1).text = str(self.title_data.get('Contract No', 'N/A'))
        table.cell(2, 0).text = 'Work Order No:'
        table.cell(2, 1).text = str(self.title_data.get('Work Order No', 'N/A'))
        
        # Add work items summary
        doc.add_heading('Work Items Summary', level=1)
        
        # Create work items table
        if not self.work_order_data.empty:
            # Add header row
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Unit'
            hdr_cells[1].text = 'Quantity executed (or supplied) since last certificate'
            hdr_cells[2].text = 'Quantity executed (or supplied) upto date as per MB'
            hdr_cells[3].text = 'S. No.'
            hdr_cells[4].text = 'Item of Work supplies'
            hdr_cells[5].text = 'Rate'
            hdr_cells[6].text = 'Upto date Amount'
            hdr_cells[7].text = 'Amount Since previous bill'
            hdr_cells[8].text = 'Remarks'
            
            # Add data rows
            for index, row in self.work_order_data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row.get('Unit', ''))
                row_cells[1].text = str(row.get('Quantity Since', ''))
                row_cells[2].text = str(row.get('Quantity Upto', ''))
                row_cells[3].text = str(row.get('Item No.', ''))
                row_cells[4].text = str(row.get('Description', ''))
                row_cells[5].text = str(row.get('Rate', ''))
                row_cells[6].text = str(row.get('Amount', ''))
                row_cells[7].text = str(row.get('Amount', ''))
                row_cells[8].text = str(row.get('Remark', ''))
        
        # Add totals section
        doc.add_heading('Totals', level=1)
        
        totals_para = doc.add_paragraph()
        totals_para.add_run(f'Grand Total: ₹{self.template_data.get("grand_total", 0):.2f}').bold = True
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_deviation_statement(self) -> bytes:
        """Generate Deviation Statement document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Deviation Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a deviation statement document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_note_sheet(self) -> bytes:
        """Generate Bill Scrutiny Sheet document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Bill Scrutiny Sheet', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a final bill scrutiny sheet document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_extra_items(self) -> bytes:
        """Generate Extra Items Statement document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extra Items Statement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is an extra items statement document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_certificate_ii(self) -> bytes:
        """Generate Certificate II document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Certificate II - Work Completion Certificate', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a work completion certificate document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    def _generate_doc_certificate_iii(self) -> bytes:
        """Generate Certificate III document in DOC format"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Certificate III - Payment Certification', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        current_date = datetime.now().strftime('%d/%m/%Y')
        date_para = doc.add_paragraph(f'Date: {current_date}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        doc.add_paragraph('This is a payment certification document.')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()